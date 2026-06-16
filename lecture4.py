"""
SmartClassNotes v3 — lecture4.py
Backend AI Engine powered by Gemini 2.5 Flash
"""
import os
import re
import json
import tempfile
import time
import google.generativeai as genai
from fpdf import FPDF
from docx import Document
import streamlit as st

GEMINI_API_KEY = st.secrets["gemini"]["api_key"] if "gemini" in st.secrets else os.environ.get("GEMINI_API_KEY")
genai.configure(api_key=GEMINI_API_KEY)

_whisper_model = None

def _get_whisper():
    global _whisper_model
    if _whisper_model is None:
        import whisper
        _whisper_model = whisper.load_model("small")
    return _whisper_model

def transcribe_audio(file_path: str) -> str:
    try: return _get_whisper().transcribe(file_path)["text"]
    except Exception as e: raise Exception(f"Audio Transcription Failed: {str(e)}")

def transcribe_youtube(youtube_url: str):
    try:
        from youtube_transcript_api import YouTubeTranscriptApi

        # Extract video ID from URL
        match = re.search(r"(?:v=|\/)([0-9A-Za-z_-]{11})", youtube_url)
        if not match:
            raise Exception("Could not extract video ID from URL.")

        video_id = match.group(1)

        # Fetch transcript
        transcript_list = YouTubeTranscriptApi.get_transcript(video_id)

        # Join all transcript text
        transcript = " ".join([entry["text"] for entry in transcript_list])

        return transcript, None
    except Exception as e:
        raise Exception(f"YouTube Transcript Failed: {str(e)}")

def _sanitize_text(text):
    if not isinstance(text, str): return str(text)
    replacements = {'\u2013': '-', '\u2014': '-', '\u2018': "'", '\u2019': "'", '\u201c': '"', '\u201d': '"', '\u2026': '...'}
    for k, v in replacements.items(): text = text.replace(k, v)
    return text.encode('latin-1', 'replace').decode('latin-1')

def _format_content_for_export(val):
    """Flattens lists and dictionaries into readable strings for full document building."""
    if isinstance(val, list): return "\n".join(f"• {v}" for v in val)
    if isinstance(val, dict): return "\n\n".join(f"{k.replace('_', ' ').title()}:\n{_format_content_for_export(v)}" for k, v in val.items())
    return str(val)

def _get_all_sections(analysis: dict):
    return [
        ("Executive Summary", analysis.get("executive_summary")),
        ("Quick Overview", analysis.get("quick_overview")),
        ("Primary Insights", analysis.get("primary_insights")),
        ("Key Points Found", analysis.get("key_points_found")),
        ("Synthesis", analysis.get("synthesis_key_points")),
        ("General Overview", analysis.get("general_overview")),
        ("Key Concepts", analysis.get("key_concepts")),
        ("Exam Preparation", analysis.get("exam_prep")),
        ("70% Length Reduction", analysis.get("reduction_70")),
        ("30% Length Reduction", analysis.get("reduction_30")),
        ("10% Length Reduction", analysis.get("reduction_10"))
    ]

def export_to_pdf(analysis: dict, out="lecture_summary.pdf"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 12, "SmartClassNotes - Complete AI Analysis", ln=True, align="C")
    pdf.ln(4)
    for title, content in _get_all_sections(analysis):
        if content:
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, title, ln=True)
            pdf.set_font("Arial", size=11)
            pdf.multi_cell(0, 7, _sanitize_text(_format_content_for_export(content)))
            pdf.ln(4)
    pdf.output(out)

def export_to_word(analysis: dict, out="lecture_summary.docx"):
    doc = Document()
    doc.add_heading("SmartClassNotes - Complete AI Analysis", 0)
    for title, content in _get_all_sections(analysis):
        if content:
            doc.add_heading(title, level=1)
            doc.add_paragraph(_sanitize_text(_format_content_for_export(content)))
    doc.save(out)

def export_to_json(analysis: dict, out="lecture_summary.json"):
    with open(out, "w", encoding="utf-8") as f:
        json.dump(analysis, f, indent=2, ensure_ascii=False)

def process_input(text=None, source_type="text", file_path=None, youtube_url=None):
    try:
        extracted_content = ""
        cleanup = []

        if source_type == "text" and text:
            extracted_content = text
        elif source_type == "file" and file_path:
            extracted_content = transcribe_audio(file_path)
        elif source_type == "youtube" and youtube_url:
            extracted_content, tmp = transcribe_youtube(youtube_url)
            if tmp:
                cleanup.append(tmp)
        else:
            return {"error": "Invalid input method."}

        if not extracted_content.strip() or len(extracted_content) < 15:
            return {"error": "No viable structural textual context could be analyzed."}

        system_prompt = '''
You are the premier AI processing core of SmartClassNotes.
Analyze the provided text and extract comprehensive academic parameters.
You MUST output ONLY a raw JSON object. Do not include markdown formatting or backticks.

Output EXACTLY this JSON structure:
{
  "executive_summary": "A high-level overview.",
  "quick_overview": "A single-sentence concise summary hook.",
  "primary_insights": "The absolute core takeaway message.",
  "reduction_percentage": "73%",
  "words_in_summary": 100,
  "key_points_found": ["Point 1", "Point 2"],
  "synthesis_key_points": ["Synthesis 1", "Synthesis 2"],
  "general_overview": "Comprehensive academic breakdown.",
  "key_concepts": ["Term 1: Definition 1", "Term 2: Definition 2"],
  "exam_prep": {
    "high_yield_insights": "Core focus areas...",
    "potential_exam_questions": ["Question 1?", "Question 2?"],
    "target_recall_terms": ["Term A", "Term B"]
  },
  "reduction_70": "Brief 30% length version of the text.",
  "reduction_30": "Medium 70% length version of the text.",
  "reduction_10": "Exhaustive 90% length retention version."
}
'''
        try:
            model = genai.GenerativeModel(
                model_name="gemini-2.5-flash",
                system_instruction=system_prompt
            )
            response = model.generate_content(
                extracted_content,
                generation_config={
                    "response_mime_type": "application/json",
                    "temperature": 0.2
                }
            )

            raw_json = response.text.strip()

            # Strip markdown fences if model hallucinates them
            if raw_json.startswith("```"):
                raw_json = re.sub(r'^```(?:json)?\n?', '', raw_json)
                raw_json = re.sub(r'\n?```$', '', raw_json)

            analysis = json.loads(raw_json)

        except Exception as e:
            return {"error": f"Gemini API Error: {str(e)}"}

        pdf_out  = "lecture_summary.pdf"
        word_out = "lecture_summary.docx"
        json_out = "lecture_summary.json"

        export_to_pdf(analysis, pdf_out)
        export_to_word(analysis, word_out)
        export_to_json(analysis, json_out)

        analysis["pdf_file"]  = pdf_out
        analysis["word_file"] = word_out
        analysis["json_file"] = json_out

        for f in cleanup:
            try:
                if f and os.path.exists(f):
                    os.remove(f)
            except Exception:
                pass

        return analysis

    except Exception as e:
        return {"error": str(e)}